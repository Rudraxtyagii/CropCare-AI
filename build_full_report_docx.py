import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import docx2pdf

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(10, 50, 25) # Deep Forest Green
        # Add bottom border or background highlight for Chapter titles
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(20, 80, 40)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(40, 100, 60)
    return p

def add_body_p(doc, text, bold_prefix=None, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = 'Arial'
        r_bold.font.size = Pt(11)
        r_bold.bold = True
        r_bold.font.color.rgb = RGBColor(30, 30, 30)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.italic = italic
    run.font.color.rgb = RGBColor(40, 40, 40)
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    
    # Create background shading for code
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(20, 20, 20)
    
    # Add light grey shading to paragraph
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F6F5"/>')
    pPr.append(shd)
    return p

def create_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Style Header Row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_background(hdr_cells[i], "145028") # Deep Green
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                
    # Style Data Rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx+1].cells
        bg_color = "F9FBF9" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)
            for paragraph in row_cells[c_idx].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx > 0 else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(40, 40, 40)
                    
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)
                
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    return table

print("Helper functions defined.")
